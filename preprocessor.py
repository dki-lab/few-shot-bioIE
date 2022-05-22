from nltk.tokenize import sent_tokenize, string_span_tokenize
import re
import os
from xml.etree import cElementTree as ET
import docx
from statistics import mean
from glob import glob
from collections import defaultdict
import json
import ipdb

def customized_sent_tokenize(text):
    sents_raw = sent_tokenize(text)
    output_sents = []
    for sent in sents_raw:
        if len(sent.split('\t')) > 1:
            output_sents.extend(sent.split('\t'))
        else:
            output_sents.append(sent)
    return output_sents


def split_sent(e1_span_s, e1_span_e, e2_span_s, e2_span_e, sent):
    # if e1 e2 not overlaping, output 5 chunks; else output 3
    pos_list = [e1_span_s, e1_span_e, e2_span_s, e2_span_e]
    if e1_span_e > e2_span_s:
        entity_s = min(e1_span_s, e2_span_s)
        entity_e = max(e1_span_e, e2_span_e)
        pos_list = [entity_s, entity_e]
    # if pos_list != sorted(pos_list):
    #     raise ValueError("Positions not in order!")
    spans = zip([0] + pos_list, pos_list + [len(sent)])
    output_chunks = []
    for (s, e) in spans:
        output_chunks.append(sent[s:e])
    return output_chunks


def extract_relation_dict(relations_file, target_labels):
    # Forming relation reference dictionary
    # {doc_id:{(e1, e2): label}}
    with open(relations_file, 'r') as f:
        relations = f.readlines()

    relation_ref_dict = {}
    for line in relations:
        doc_id, label, _, _, e1, e2 = line.rstrip().split('\t')
        e1_id = e1.split(':')[1]
        e2_id = e2.split(':')[1]
        if doc_id not in relation_ref_dict:
            relation_ref_dict[doc_id] = {}
        label = label if label in target_labels else "false"
        relation_ref_dict[doc_id][(e1_id, e2_id)] = label
    return relation_ref_dict


def extract_entity_dict(entities_file):
    # entity span refer
    # {doc_id:[[e_id, type, span_s, span_e, content]]}
    with open(entities_file, 'r') as f:
        entities = f.readlines()
    entity_span_dict = {}
    for line in entities:
        doc_id, e_id, type, span_s, span_e, content = line.rstrip().split('\t')
        if doc_id not in entity_span_dict:
            entity_span_dict[doc_id] = []
        # Ignoring the suffixe
        type = type.split('-')[0]
        entity_span_dict[doc_id].append(
            [e_id, type, int(span_s), int(span_e), content])
    return entity_span_dict


def reformat_data(abstract_file, relation_ref_dict, entity_span_dict):
    # Traversing abstract, and finding candidates with exact one chem
    # and one gene
    with open(abstract_file, 'r') as f:
        abstract_data = f.readlines()

    processed_data_mask = []
    processed_data_mask.append(
        f"id\torigin_sent\tdrug\tgene\tsent\tlabel\n")

    for line in abstract_data:
        doc_id, text = line.split('\t', 1)
        sents = customized_sent_tokenize(text)
        entity_candidates = entity_span_dict[doc_id]
        prev_span_end = 0
        for sent in sents:
            # Extacting span of cur sent.
            sent_span_s = text.find(sent, prev_span_end)
            sent_span_e = sent_span_s + len(sent)
            prev_span_end = sent_span_e
            chem_list = []
            gene_list = []
            for entity_candidate in entity_candidates:
                e_id, type, entity_span_s, entity_span_e, content = \
                    entity_candidate
                if entity_span_s >= sent_span_s and entity_span_e \
                        <= sent_span_e:
                    if "CHEM" in type:
                        chem_list.append(entity_candidate)
                    else:
                        gene_list.append(entity_candidate)
            if len(chem_list) == 0 or len(gene_list) == 0:
                continue

            # Preparing data with appending method
            for chem_candidate in chem_list:
                for gene_candidate in gene_list:

                    gene_id, gene_type, gene_span_s, gene_span_e, gene_content = gene_candidate
                    chem_id, chem_type, chem_span_s, chem_span_e, chem_content = chem_candidate

                    # Denoting the first entity entity 1.
                    if chem_candidate[2] < gene_candidate[2]:
                        e1_candidate, e2_candidate = \
                            chem_candidate, gene_candidate
                    else:
                        e2_candidate, e1_candidate = \
                            chem_candidate, gene_candidate

                    e1_id, e1_type, e1_span_s, e1_span_e, e1_content = \
                        e1_candidate
                    e2_id, e2_type, e2_span_s, e2_span_e, e2_content = \
                        e2_candidate

                    label = "false"

                    processed_doc_id = f"{doc_id}.{e1_id}.{e2_id}"
                    if doc_id in relation_ref_dict:
                        if (e1_id, e2_id) in relation_ref_dict[doc_id]:
                            label = relation_ref_dict[doc_id][(e1_id, e2_id)]
                        elif (e2_id, e1_id) in relation_ref_dict[doc_id]:
                            label = relation_ref_dict[doc_id][(e2_id, e1_id)]

                    e1_span_s, e1_span_e, e2_span_s, e2_span_e = \
                        e1_span_s - sent_span_s, e1_span_e - sent_span_s, \
                        e2_span_s - sent_span_s, e2_span_e - sent_span_s
                    # split sent into chunks
                    chunks = split_sent(
                        e1_span_s, e1_span_e, e2_span_s, e2_span_e, sent)
                    if len(chunks) == 5:
                        chunk1, chunk2_e1, chunk3, chunk4_e2, chunk5 = chunks
                        processed_sent_mask = \
                            f"{chunk1}{'ENT1'}{chunk3}{'ENT2'}{chunk5}"
                    else:
                        chunk1, chunk2_entity, chunk3 = chunks
                        entity_type = "ENT1-ENT2"
                        processed_sent_mask = \
                            f"{chunk1}{entity_type}{chunk3}"

                    # Forming sent using mask method
                    processed_data_mask.append(
                        f"{processed_doc_id}\t{clean_sent(sent)}\t{clean_sent(chem_content)}\t{clean_sent(gene_content)}\t{clean_sent(processed_sent_mask)}\t{label}\n")
    return processed_data_mask


def clean_sent(sent):
    special_chars = ['\n', '\t', '\r']
    for special_char in special_chars:
        sent = sent.replace(special_char, ' ')
    return sent


def extract_entity_dict_DDI(entity_candidates):
    candidate_ref_dict = {}
    for entity_candidate in entity_candidates:
        candidate_ref_dict[entity_candidate.get(
            'id')] = entity_candidate.attrib
    return candidate_ref_dict


def extract_span_DDI(span_str):
    candidates = re.findall("\d+", span_str)
    # When multiple spans occurs, only taking the very first and last positions
    # Ending position offsets by 1
    span_s, span_e = int(candidates[0]), int(candidates[-1]) + 1

    return [span_s, span_e]


def dump_processed_data(output_dir, data_type, processed_data_mask):
    processed_data_mask_file = os.path.join(output_dir, f"{data_type}.tsv")
    os.makedirs(os.path.dirname(processed_data_mask_file), exist_ok=True)

    with open(processed_data_mask_file, 'w') as f:
        f.writelines(processed_data_mask)


def prepare_chemprot_data(root_dir, output_dir):
    data_types = ['train', 'dev', 'test']
    target_labels = ["CPR:3", "CPR:4", "CPR:5", "CPR:6", "CPR:9"]

    for data_type in data_types:
        if data_type == 'train':
            data_path = os.path.join(
                root_dir, 'chemprot_training/chemprot_training/')
            file_name_prefix = "chemprot_training_"
            file_name_affixe = ""
        elif data_type == "dev":
            data_path = os.path.join(
                root_dir, 'chemprot_development/chemprot_development/')
            file_name_prefix = "chemprot_development_"
            file_name_affixe = ""
        else:
            data_path = os.path.join(
                root_dir, 'chemprot_test_gs/chemprot_test_gs/')
            file_name_prefix = "chemprot_test_"
            file_name_affixe = "_gs"

        relations_file = os.path.join(
            data_path, f"{file_name_prefix}relations{file_name_affixe}.tsv")
        entities_file = os.path.join(
            data_path, f"{file_name_prefix}entities{file_name_affixe}.tsv")
        abstract_file = os.path.join(
            data_path, f"{file_name_prefix}abstracts{file_name_affixe}.tsv")

        relation_ref_dict = extract_relation_dict(
            relations_file, target_labels)
        entity_span_dict = extract_entity_dict(entities_file)
        processed_data_mask = reformat_data(
            abstract_file, relation_ref_dict, entity_span_dict)
        # Dumping data.
        dump_processed_data(output_dir, data_type, processed_data_mask)


def prepare_DDI_data(root_dir, output_dir):
    data_types = ["train", "dev", "test"]
    # prepare train/dev/test file names

    for data_type in data_types:
        processed_data_appending = []
        processed_data_boundary = []
        processed_data_mask = []
        processed_data_mask.append(
            f"id\torigin_sent\tdrug1\tdrug2\tsent\tlabel\n")
        # load file names
        with open(f"indexing/DDI/{data_type}_files.tsv", 'r') as f:
            data_file_names = f.readlines()
        data_file_paths = [root_dir+x.strip() for x in data_file_names]
        print(data_type+f":{len(data_file_paths)}")
        for data_file_path in data_file_paths:
            # data_file = os.path.join(data_dir, data_file_path)
            tree = ET.parse(data_file_path)
            root = tree.getroot()
            for sent in list(root):
                if sent.find('pair') is None:
                    continue
                entity_candidates = sent.findall('entity')
                # candidate_ref_dict = {}
                # for entity_candidate in entity_candidates:
                #     candidate_ref_dict[entity_candidate.get('id')] = entity_candidate.attrib
                candidate_ref_dict = extract_entity_dict_DDI(entity_candidates)
                text = sent.attrib.get('text')
                entity_candidates = sent.findall('entity')
                candidate_ref_dict = {}
                for entity_candidate in entity_candidates:
                    candidate_ref_dict[entity_candidate.get(
                        'id')] = entity_candidate.attrib
                pairs = sent.findall('pair')
                for pair in pairs:
                    pair_id, e1_id, e2_id, label = pair.attrib['id'], pair.attrib[
                        'e1'], pair.attrib['e2'], pair.attrib['ddi']
                    if label == 'true':
                        label = pair.attrib['type']
                    label = f"DDI-{label}"
                    e1 = candidate_ref_dict[e1_id]
                    e2 = candidate_ref_dict[e2_id]

                    # Ensuring e1 is the first entity.
                    if extract_span_DDI(e1['charOffset'])[0] > extract_span_DDI(e2['charOffset'])[0]:
                        e1, e2 = e2, e1
                    e1_content = e1['text']
                    e2_content = e2['text']
                    e1_span_s, e1_span_e = extract_span_DDI(e1['charOffset'])
                    e2_span_s, e2_span_e = extract_span_DDI(e2['charOffset'])

                    # Forming sent using appending method
                    append_entity_content = f"{e1_content},{e2_content}"
                    # processed_sent_appending = f"{text} {append_entity_content}"
                    processed_sent_appending = text

                    # processed_line_appending = f"{pair_id}\t{processed_sent_appending}\tDDI-{label}\n"

                    # Forming sent using boundary tokens method
                    chunks = split_sent(
                        e1_span_s, e1_span_e, e2_span_s, e2_span_e, text)
                    if len(chunks) == 5:
                        chunk1, chunk2_e1, chunk3, chunk4_e2, chunk5 = chunks
                        # processed_sent_boundary = f"{chunk1}<e1>{chunk2_e1}</e1>{chunk3}<e2>{chunk4_e2}</e2>{chunk5}"
                        processed_sent_boundary = f"{chunk1}{chunk2_e1}{chunk3}{chunk4_e2}{chunk5}"
                        processed_sent_mask = f"{chunk1}ENT1{chunk3}ENT2{chunk5}"
                    else:
                        chunk1, chunk2_entity, chunk3 = chunks
                        # processed_sent_boundary = f"{chunk1}<e1>{chunk2_entity}</e1>{chunk3}"
                        processed_sent_boundary = f"{chunk1}@{chunk2_entity}${chunk3}"

                        entity_type = "ENT1-ENT2"
                        processed_sent_mask = f"{chunk1}{entity_type}{chunk3}"
                        # print(e1_span_s, e1_span_e, e2_span_s, e2_span_e)

                    processed_data_appending.append(
                        f"{pair_id}\t{clean_sent(processed_sent_appending)}\t{label}\n")
                    # Forming sent using boundary tokens method
                    processed_data_boundary.append(
                        f"{pair_id}\t{clean_sent(processed_sent_boundary)}\t{label}\n")

                    # Forming sent using mask method
                    processed_data_mask.append(
                        f"{pair_id}\t{clean_sent(text)}\t{clean_sent(e1_content)}\t{clean_sent(e2_content)}\t{clean_sent(processed_sent_mask)}\t{label}\n")
        dump_processed_data(output_dir, data_type, processed_data_mask)


def preprocess_ner_data(input_path, output_path):
    data_types = ["train.tsv", "dev.tsv", "devel.tsv", 'test.tsv']
    data_type_mapping = \
        {"train.tsv": "train.conll.csv", "devel.tsv": "dev.conll.csv", "dev.tsv": "dev.conll.csv", "test.tsv": "test.conll.csv"}
    os.makedirs(output_path, exist_ok=True)
    for data_type in data_types:
        input_file_path = os.path.join(input_path, data_type)
        output_file_path = os.path.join(
            output_path, data_type_mapping[data_type])

        try:
            with open(input_file_path, "r") as fin:
                data_in = fin.readlines()
        except:
            print('{} {}'.format(input_path, data_type))
            continue

        fout = open(output_file_path, "w")
        past_line = '\n'
        fout.write('tokens\tner_tags\n')
        for line in data_in:
            if not line.split() and past_line != '\n' and not(past_line.startswith("-DOCSTART-")):
                fout.write(line)
            elif not(line.startswith("-DOCSTART-")) and line != '\n':
                token, label = line.split('\t')

                fout.write(f"{token}\t{label[0]}\n")

            past_line = line

        fin.close()
        fout.close()


def prepare_jnlpba_data(input_path, output_path):
    train_file_path = os.path.join(input_path, "Genia4ERtask2.iob2")
    test_file_path = os.path.join(input_path, "Genia4EReval2.iob2")
    with open(train_file_path, 'r') as f:
        train_raw = f.readlines()
    with open(test_file_path, 'r') as f:
        test = f.readlines()

    def change_doc_start(input_data):
        for i in range(len(input_data)):
            if input_data[i].startswith("###MEDLINE"):
                input_data[i] = input_data[i].replace(
                    "###MEDLINE:", "-DOCSTART- ")
        return input_data
    train_raw = change_doc_start(train_raw)
    test = change_doc_start(test)

    # To ensure fair comparision, we adopted the same train/dev split as MTL.
    split_idx = 465497
    train, dev = train_raw[:split_idx], train_raw[split_idx:]

    def dump_jnlpba_data(data, output_path):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w') as writer:
            writer.writelines(data)
    dump_jnlpba_data(train, os.path.join(output_path, 'train.tsv'))
    dump_jnlpba_data(dev, os.path.join(output_path, 'dev.tsv'))
    dump_jnlpba_data(test, os.path.join(output_path, 'test.tsv'))
    preprocess_ner_data(output_path, output_path)


def clean_sent_biosses(sent):
    special_chars = ['\n', '\t']
    for special_char in special_chars:
        sent = sent.replace(special_char, ' ')
    return sent.strip()


def extract_doc_table(table):
    """
    return list of dict.
    """
    # Data will be a list of rows represented as dictionaries
    # containing each row's data.
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data


def prepare_biosses_data(input_path, output_path):
    pair_file_name = os.path.join(input_path, 'Annotation-Pairs.docx')
    score_file_name = os.path.join(input_path, 'Annotator-Scores.docx')
    score_table = docx.Document(score_file_name).tables[0]
    pair_table = docx.Document(pair_file_name).tables[0]
    score_data = extract_doc_table(score_table)
    pair_data = extract_doc_table(pair_table)
    if len(score_data) != len(pair_data):
        raise ValueError("Pair file and score file mismaches!")
    processed_data = []
    processed_data_dict = {}
    for idx in range(len(pair_data)):
        pair_raw = pair_data[idx]
        score_raw = score_data[idx]
        pair_id, sent1, sent2 = pair_raw[''].strip(
        ), pair_raw['Sentence 1'], pair_raw['Sentence 2']
        score_id, scores = score_raw['Sentence No'], [score_raw['Annotator A'], score_raw['Annotator B'],
                                                      score_raw['Annotator C'], score_raw['Annotator D'],
                                                      score_raw['Annotator E']]
        if pair_id != score_id:
            raise ValueError("Id Mismatches.")
        score = float(mean([int(x) for x in scores]))
        processed_data.append([pair_id, clean_sent_biosses(
            sent1), clean_sent_biosses(sent2), score])
        processed_data_dict[pair_id] = [clean_sent_biosses(
            sent1), clean_sent_biosses(sent2), score]

    # There is no train/dev/test split in official data.
    # FOr fair comparison, we extracted indexes as below from BLUE.
    train_indexes = ['78', '45', '35', '50', '27', '13', '87', '1', '58', '99',
                     '55', '74', '66', '39', '44', '18', '84', '76', '19', '10',
                     '75', '46', '15', '86', '60', '14', '51', '79', '29', '34',
                     '94', '28', '62', '42', '21', '30', '11', '53', '6', '12',
                     '26', '48', '31', '32', '77', '37', '95', '85', '36', '56',
                     '43', '61', '16', '5', '67', '65', '54', '3', '73', '98',
                     '17', '4', '92', '93']
    dev_indexes = ['88', '82', '8', '63', '47', '68', '40',
                   '90', '100', '24', '41', '91', '80', '9', '72', '2']
    test_indexes = ['59', '96', '70', '22', '81', '38', '57', '23', '33',
                    '89', '69', '49', '7', '71', '97', '25', '83', '64', '52', '20']

    def dump_biosses_data(indexes, output_path, processed_data_dict):
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w') as writer:
            writer.write("pair_id\tsent1\tsent2\tscore\n")
            for idx, pair_id in enumerate(indexes):
                sent1, sent2, score = processed_data_dict[pair_id]
                line = f"{pair_id}\t{sent1}\t{sent2}\t{score}\n"
                writer.write(line)
    dump_biosses_data(train_indexes, os.path.join(
        output_path, "train.tsv"), processed_data_dict)
    dump_biosses_data(dev_indexes, os.path.join(
        output_path, "dev.tsv"), processed_data_dict)
    dump_biosses_data(test_indexes, os.path.join(
        output_path, "test.tsv"), processed_data_dict)


def fname_to_pmid(fname):
    pmid = os.path.splitext(os.path.basename(fname))[0].split('.')[0]
    return pmid


def preprocess_pico_data(input_path, output_path):
    # script based on https://github.com/bepnye/EBM-NLP/blob/master/models/lstm-crf/build_data.py
    # We sorted the files in alphabet order to ensure consistent train dev split.

    # Loading indexing extracted from scibert
    data_indexing = defaultdict(list)
    for batch in ["train", "dev", "test"]:
        with open(f"indexing/ebmnlp/{batch}_pmid.tsv", 'r') as f:
            data_indexing[batch] = f.readline().split(',')
    ebm_nlp = input_path

    id_to_tokens = {}
    id_to_pos = {}
    PIO = ['participants', 'interventions', 'outcomes']
    PHASES = ['starting_spans']
    label_mapping = {'p': 'I-PAR', 'i': 'I-INT', 'o': 'I-OUT'}
    token_fnames = glob('%s/documents/*.tokens' % ebm_nlp)
    for fname in token_fnames:
        pmid = fname_to_pmid(fname)
        tokens = open(fname).read().split()
        tags = open(fname.replace('tokens', 'pos')).read().split()
        id_to_tokens[pmid] = tokens
        id_to_pos[pmid] = tags

    batch_to_labels = {}
    for phase in PHASES:
        batch_to_labels[phase] = {}
        for pio in PIO:
            batch_to_labels[phase][pio] = {}
            print('Reading files for %s %s' % (phase, pio))
            for fdir in ['train', 'test/gold']:
                batch = fdir.split('/')[0]
                batch_to_labels[phase][pio][batch] = dict()
                ann_fnames = glob(
                    '%s/annotations/aggregated/%s/%s/%s/*.ann' % (ebm_nlp, phase, pio, fdir))
                for fname in ann_fnames:
                    pmid = fname_to_pmid(fname)
                    batch_to_labels[phase][pio][batch][pmid] = open(
                        fname).read().split()

    batch_groups = [('p1_all', ['starting_spans'], [
                     'participants', 'interventions', 'outcomes'])]

    for group_name, phases, pio in batch_groups:

        id_to_labels_list = defaultdict(list)
        batch_to_ids = defaultdict(set)
        for phase in phases:
            for e in pio:
                print('Collecting anns from %s %s' % (phase, e))
                for batch, batch_labels in batch_to_labels[phase][e].items():
                    print('\t%d ids for %s' % (len(batch_labels), batch))
                    batch_to_ids[batch].update(batch_labels.keys())
                    for pmid, labels in batch_labels.items():
                        labels = ['%s_%s' % (l, e[0]) for l in labels]
                        id_to_labels_list[pmid].append(labels)

        for batch, ids in batch_to_ids.items():
            print('Found %d ids for %s' % (len(ids), batch))

        # loading indexes from scibert
        batch_to_ids = data_indexing
        for batch, ids in batch_to_ids.items():
            os.makedirs(output_path, exist_ok=True)

            fout = open(os.path.join(output_path, f"{batch}.tsv"), 'w')
            for pmid in ids:
                fout.write(f"-DOCSTART- ({pmid})\n\n")
                tokens = id_to_tokens[pmid]
                poss = id_to_pos[pmid]
                per_token_labels = zip(*id_to_labels_list[pmid])
                for i, (token, pos, labels) in enumerate(zip(tokens, poss, per_token_labels)):
                    final_label = 'O'
                    for l in labels:
                        if l[0] != '0':
                            # final_label = label_mapping[l]
                            final_label = label_mapping[l.split('_')[-1]]
                    fout.write(f"{token}\t{final_label}\n")
                    # '%s %s %s\n' % (token, pos, final_label))
                    if token == '.':
                        fout.write('\n')


def preprocess_hoc_data(input_path, output_path):

    label_dir = os.path.join(input_path, "labels")
    text_dir = os.path.join(input_path, "text")

    label_file_names = sorted(
        [x for x in os.listdir(label_dir) if x.endswith('.txt')])
    LABELS = ['activating invasion and metastasis', 'avoiding immune destruction',
              'cellular energetics', 'enabling replicative immortality', 'evading growth suppressors',
              'genomic instability and mutation', 'inducing angiogenesis', 'resisting cell death',
              'sustaining proliferative signaling', 'tumor promoting inflammation']

    def format_hoc_instances(label_file_names):
        processed_data = []

        for label_file_name in label_file_names:
            with open(os.path.join(label_dir, label_file_name), 'r') as f:
                label_data = f.readlines()
            with open(os.path.join(text_dir, label_file_name), 'r') as f:
                text_data = f.readlines()
            # Ignoring the leading space.
            labels_raw = label_data[0].split('<')[1:]
            # Processing each sentence label list pair.
            for i in range(len(text_data)):
                text = text_data[i]
                label_raw = labels_raw[i]
                label_list = [0]*10
                # Transcribe label to index
                for lab in label_raw.split('AND'):
                    if lab not in [' ']:
                        top_level_label = lab.split('--')[0].lower().lstrip()
                        if 'null' in top_level_label:
                            continue
                        label_idx = LABELS.index(top_level_label)
                        if label_idx < 0:
                            raise ValueError("label not in list")
                        label_list[label_idx] = 1
                label_final = ','.join(
                    [f"{idx}_{x}" for idx, x in enumerate(label_list)])
                line = f"{label_final}\t{text.strip()}\t{label_file_name.split('.')[0]}_s{i}\n"
                processed_data.append(line)
        return processed_data
    data_indexing = defaultdict(list)
    output_data = {}
    for batch in ["train", "dev", "test"]:
        with open(f"indexing/HoC/{batch}_pmid.tsv", 'r') as f:
            data_indexing[batch] = f.readline().split(',')

        output_data[batch] = format_hoc_instances(data_indexing[batch])
    os.makedirs(output_path, exist_ok=True)
    for key, val in output_data.items():
        with open(os.path.join(output_path, f"{key}.tsv"), 'w') as fout:
            fout.write("labels\tsentence\tindex\n")
            fout.writelines(val)


def preprocess_bioasq_corpus(input_path, output_path):
    train_data_path = os.path.join(input_path, "BioASQ-training7b/trainining7b.json")
    test_dir = os.path.join(input_path, "Task7BGoldenEnriched")

    data_dict = {}
    with open(train_data_path, 'r') as f:
        data = json.load(f)
    for d in data['questions']:
        type = d['type']
        if type != "yesno":
            continue
        question, label, doc_id = d['body'].replace(
            '\n', ' '), d['exact_answer'], d['id']
        snippets = d['snippets']
        snippets_text = [snippet['text'].replace(
            '\n', ' ') for snippet in snippets]
        passage = " ".join(snippets_text)
        text_line = f"{doc_id}\t{question}\t{passage}\t{label}\n"
        data_dict[doc_id] = text_line

    for test_file in os.listdir(test_dir):
        test_data_path = os.path.join(test_dir, test_file)
        print(test_data_path)
        with open(test_data_path, 'r') as f:
            data = json.load(f)
        for d in data['questions']:
            type = d['type']
            if type != "yesno":
                continue
            question, label, doc_id = d['body'].replace(
                '\n', ' '), d['exact_answer'], d['id']
            snippets = d['snippets']
            snippets_text = [snippet['text'].replace(
                '\n', ' ') for snippet in snippets]
            passage = " ".join(snippets_text)
            text_line = f"{doc_id}\t{question}\t{passage}\t{label}\n"
            data_dict[doc_id] = text_line

    # load ids
    batches = ['train', 'dev', 'test']
    for batch in batches:
        with open(f"indexing/BioASQ/{batch}_id.tsv", 'r') as f:
            ids = f.readline().split(',')
        lines = []
        for doc_id in ids:
            lines.append(data_dict[doc_id])
        os.makedirs(output_path, exist_ok=True)
        with open(os.path.join(output_path, f"{batch}.tsv"), 'w') as f:
            f.writelines(lines)


def preprocess_GAD_corpus(input_path, output_path):
    train_data_path = os.path.join(input_path, "1/train.tsv")
    with open(train_data_path, 'r') as f:
        train_data = f.readlines()
    data = {}
    data['train'], data['dev'] = train_data[:4261], train_data[4261:]
    os.makedirs(output_path, exist_ok=True)
    for batch in ['train', 'dev']:
        with open (os.path.join(output_path, f'{batch}.tsv') , 'w') as f:
            f.write("index\tsentence\tlabel\n")
            for idx, line in enumerate(data[batch]):
                f.write(f"{idx}\t{line}")
    with open(os.path.join(input_path, "1/test.tsv"), 'r') as f:
        test_data = f.readlines()
    with open(os.path.join(output_path, 'test.tsv') , 'w') as f:
        f.writelines(test_data)


def main():
    print("Processing DDI corpus...")
    prepare_DDI_data("raw_data/DDI/DDIextraction_2013/", "data/DDI")
  
    print("Processing NER data...")
    # NER
    input_path = "raw_data/MTL-Bioinformatics-2016/data/"
    output_path = "data/"
    preprocess_ner_data(input_path+"BC5CDR-disease-IOB",
                        output_path+"BC5CDR-disease")
    preprocess_ner_data(input_path+"BC5CDR-chem-IOB",
                        output_path+"BC5CDR-chem")
    preprocess_ner_data(input_path+"BC2GM-IOB", output_path+"BC2GM")
    preprocess_ner_data(input_path+"NCBI-disease-IOB",
                        output_path+"NCBI-disease")
    prepare_jnlpba_data("raw_data/JNLPBA/",  output_path+"JNLPBA")

    print("Processing GAD data...")
    preprocess_GAD_corpus("raw_data/GAD", "data/GAD")

    if os.path.exists('raw_data/ChemProt_Corpus'):
        print("Processing chemprot corpus...")
        prepare_chemprot_data('raw_data/ChemProt_Corpus', 'data/chemprot')
if __name__ == "__main__":
    main()
